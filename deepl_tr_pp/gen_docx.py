"""Generate docx from two lists/texts."""
# pylint: disable=duplicate-code

from typing import (
    List,
    Union,
)

from pathlib import Path
from itertools import zip_longest

# from textwrap import shorten

from docx import Document

# from docx.shared import Pt
# from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

from logzero import logger


# fmt: off
def gen_docx(
        src: Union[str, List[str]],
        tgt: Union[str, List[str]],
        template: str = "templ_dual.docx",
) -> Document:
    # fmt: on
    """Generate docx from two lists/texts.

    tempplate default to templ_dual.docx in the current directory if pesent
    """
    if isinstance(src, str):
        src = [elm.strip() for elm in src.splitlines() if elm.strip()]
    if isinstance(tgt, str):
        tgt = [elm.strip() for elm in tgt.splitlines() if elm.strip()]

    if not len(src) == len(tgt):
        logger.warning(" lent(src) %s and len(tgt) %s not match", len(src), len(tgt))
        logger.warning("There appears to be some problem.")
        logger.warning("We proceed nevertheless.")

    _ = Path(template).expanduser().resolve()
    if _.exists():
        document = Document(_)
        logger.info("Using template %s", template)
    else:
        logger.info(" %s not present, no template file used", template)
        logger.info("(A template file dictates fonts, line spacing, margins, etc.)")
        document = Document()

    # def add_para(document, elm: str, color: bool = True, highlight: bool = False):
    def add_para(document, elm: str, color: bool = False, highlight: bool = True):
        paragraph = document.add_paragraph()

        # remove leading and trailing spaces
        try:
            elm = elm.strip()
        except Exception as exc:
            logger.error(exc)

        run = paragraph.add_run(elm)
        font = run.font
        if highlight:
            # font.highlight_color = WD_COLOR_INDEX.GRAY_25  # pylint: disable=E1101
            # font.highlight_color = WD_COLOR_INDEX.WHITE  # pylint: disable=E1101  8
            font.highlight_color = WD_COLOR_INDEX.YELLOW  # pylint: disable=E1101  8

            # https://www.rapidtables.com/web/color/Yellow_Color.html
            # lightyellow	#FFFFE0	rgb(255,255,224) 
            # yellow	#FFFF00	rgb(255,255,0)
            # font.highlight_color = RGBColor(0xFF, 0xFF, 0xE0)  # does no work pylint: disable=E1101
        if color:
            font.color.rgb = RGBColor(0xff, 0xff, 0xe0)  # noqa

    for elm0, elm1 in zip_longest(src, tgt, fillvalue=""):
        add_para(document, elm0, False, True)  # 
        add_para(document, elm1, False, False)  # 

    return document
